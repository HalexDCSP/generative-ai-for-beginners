import unittest
from io import BytesIO

from docx import Document

from ata_reuniao_core import (
    compose_preview_text,
    extract_json_payload,
    extract_template_blueprint,
    fill_template_by_spec,
    parse_preview_text,
    preview_to_payload,
)


def sample_preview_text() -> str:
    return compose_preview_text(
        {
            "cliente": "Cliente Exemplo",
            "nome_codigo_projeto": "Projeto XPTO - PRJ001",
            "data_reuniao": "10-Mar-2026",
            "horas": "09:30",
            "data_documento": "10-Mar-2026",
            "convocados": ["Ana Souza | Cliente Exemplo - TI | Não"],
            "pauta": ["Alinhamento inicial", "Planejamento da entrega"],
            "definicoes": ["Cronograma aprovado"],
            "pendencias": ["Enviar acessos | Ana Souza | 12-Mar-2026"],
            "plano_acao": ["Realizar kickoff técnico | Equipe Senior | 15-Mar-2026"],
        }
    )


class AtaReuniaoCoreTests(unittest.TestCase):
    def test_extract_json_payload_accepts_fenced_json(self) -> None:
        payload = extract_json_payload("```json\n{\"ok\": true}\n```")
        self.assertTrue(payload["ok"])

    def test_preview_roundtrip_preserves_main_fields(self) -> None:
        preview = sample_preview_text()
        parsed = parse_preview_text(preview)
        self.assertEqual(parsed["cliente"], "Cliente Exemplo")
        self.assertEqual(parsed["horas"], "09:30")
        self.assertIn("Alinhamento inicial;", parsed["pauta"])

    def test_extract_template_blueprint_reads_paragraphs_and_cells(self) -> None:
        document = Document()
        document.add_paragraph("Cliente: [Inserir nome do cliente]")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Data"
        table.cell(0, 1).text = "[DD-MMM-AAAA]"
        buffer = BytesIO()
        document.save(buffer)
        blueprint = extract_template_blueprint(buffer.getvalue())

        self.assertEqual(blueprint["paragraphs"][0]["id"], "P001")
        self.assertEqual(blueprint["cells"][1]["id"], "T001R001C002")
        self.assertIn("Inserir nome do cliente", blueprint["placeholders"])

    def test_fill_template_by_spec_populates_fixed_and_dynamic_tables(self) -> None:
        document = Document()
        # Tabela 0 - Informações do Projeto
        t0 = document.add_table(rows=5, cols=3)
        t0.cell(1, 1).text = "[Inserir nome do cliente]"
        t0.cell(2, 1).text = "[Inserir nome e código do projeto]"
        t0.cell(3, 1).text = "[DD-MMM-AAAA]"
        t0.cell(3, 2).text = "Horas: [HH:MM]"
        t0.cell(4, 1).text = "[DD-MMM-AAAA]"

        # Tabela 1 - Convocados
        t1 = document.add_table(rows=2, cols=3)
        t1.cell(1, 0).text = "[Nome do participante]"
        t1.cell(1, 1).text = "[Setor/Departamento]"
        t1.cell(1, 2).text = "[Sim/Não]"

        # Tabelas 2 e 3
        t2 = document.add_table(rows=2, cols=1)
        t2.cell(1, 0).text = "[Inserir itens]"
        t3 = document.add_table(rows=2, cols=1)
        t3.cell(1, 0).text = "[Inserir definições]"

        # Tabela 4 - Pendências
        t4 = document.add_table(rows=2, cols=3)
        t4.cell(1, 0).text = "[Descrição da pendência]"
        t4.cell(1, 1).text = "[Responsável]"
        t4.cell(1, 2).text = "[DD-MMM-AAAA]"

        # Tabela 5 - Plano de Ação
        t5 = document.add_table(rows=2, cols=3)
        t5.cell(1, 0).text = "[Descrição da ação]"
        t5.cell(1, 1).text = "[Responsável]"
        t5.cell(1, 2).text = "[DD-MMM-AAAA]"

        buffer = BytesIO()
        document.save(buffer)

        preview = sample_preview_text().replace(
            "Pauta:\n",
            "- Bruno Lima | Cliente Exemplo - Operações | Não\n\nPauta:\n",
        )
        filled_document, report = fill_template_by_spec(buffer.getvalue(), preview)

        self.assertEqual(filled_document.tables[0].cell(1, 1).text, "Cliente Exemplo")
        self.assertEqual(filled_document.tables[0].cell(3, 2).text, "Horas: 09:30")
        self.assertEqual(len(filled_document.tables[1].rows), 3)
        self.assertEqual(filled_document.tables[1].cell(2, 0).text, "Bruno Lima")
        self.assertGreaterEqual(len(filled_document.tables[4].rows), 2)
        self.assertGreaterEqual(len(filled_document.tables[5].rows), 2)
        self.assertEqual(report["linhas_convocados"], 2)

        cliente_run = filled_document.tables[0].cell(1, 1).paragraphs[0].runs[0]
        self.assertEqual(cliente_run.font.name, "Calibri")
        self.assertEqual(cliente_run.font.bold, False)
        self.assertEqual(cliente_run.font.size.pt, 11)

    def test_preview_to_payload_parses_lists(self) -> None:
        payload = preview_to_payload(sample_preview_text())
        self.assertEqual(payload["convocados"][0]["nome"], "Ana Souza")
        self.assertEqual(payload["pendencias"][0]["prazo"], "12-Mar-2026")


if __name__ == "__main__":
    unittest.main()