"""Funções centrais para análise e preenchimento de atas de reunião."""

from __future__ import annotations

import copy
import json
import re
import unicodedata
from copy import deepcopy
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


PROMPT_1_PREVIA = """Analise a transcrição completa desta reunião e extraia TODAS as informações listadas abaixo, seguindo rigorosamente o formato e as regras especificadas.

REGRAS DE IDENTIFICAÇÃO DO CLIENTE E PROJETO:
1. Identifique o nome do cliente a partir de qualquer uma destas fontes, em ordem de prioridade:
   a) Nome do cliente mencionado explicitamente na conversa
   b) Título da reunião no Teams (geralmente contém o nome do cliente ou projeto)
   c) Nome da empresa dos participantes externos (que não são da Senior)
   d) Contexto da conversa: referências a sistemas, ambientes ou operações do cliente
   Se mesmo assim não for possível, use o nome da empresa que aparece nos participantes do Teams que NÃO sejam da Senior.
2. Identifique o nome e código do projeto da mesma forma: menção direta, título da reunião, ou contexto (ex: "projeto WMS", "implantação ERP", nome do sistema discutido).

REGRAS DE IDENTIFICAÇÃO DOS PARTICIPANTES:
3. Use os NOMES COMPLETOS conforme aparecem na lista de participantes do Teams, não apenas o primeiro nome.
4. Para o Setor/Departamento de cada participante, use esta hierarquia:
   a) Cargo ou departamento explicitamente mencionado na conversa
   b) Informação de cargo/empresa visível no perfil do Teams do participante
   c) Inferência pelo papel na conversa: quem discute questões técnicas de banco de dados é provavelmente DBA ou TI; quem coordena tarefas e prazos é provavelmente GP ou Consultor; quem fala sobre operação do CD é provavelmente Operações do cliente; quem trata de configuração de sistema é provavelmente Consultor Funcional ou Técnico
   d) Se a pessoa é da Senior, escreva "Senior - [cargo/função inferida]"
   e) Se a pessoa é do cliente, escreva "[Nome do cliente] - [cargo/função inferida]"
   f) Somente se não houver nenhuma pista, use "Não identificado"

REGRAS GERAIS DE EXTRAÇÃO:
5. Datas DEVEM estar no formato DD-MMM-AAAA com meses SEMPRE abreviados em português: Jan, Fev, Mar, Abr, Mai, Jun, Jul, Ago, Set, Out, Nov, Dez.
6. Horas DEVEM estar no formato HH:MM (24h).
7. A Data de criação do Documento é sempre igual à Data da Reunião.
8. Ausência: marque "Sim" APENAS se o participante foi explicitamente mencionado como ausente ou não compareceu. Caso contrário, marque "Não".
9. Pauta: cada item DEVE terminar com ";" (ponto e vírgula). Identifique os temas principais discutidos.
10. Definições: cada item em uma linha separada, SEM ";" ao final. Definições são decisões tomadas, conclusões alcançadas e fatos estabelecidos durante a reunião. NÃO deixe linhas em branco entre os itens nem após o último item.
11. Pendências: tarefas que alguém precisa entregar ou resolver, com responsável e prazo. Se o prazo não foi definido, use "A definir".
12. Plano de Ação: ações futuras planejadas ou decididas, com responsável e prazo. Se o prazo não foi definido, use "A definir".
13. NÃO invente informações. Extraia apenas o que está na transcrição ou no contexto do Teams.

FORMATO DE SAÍDA (reproduza exatamente esta estrutura, sem formatação adicional):

Cliente: [nome do cliente]
Nome e Código do projeto: [nome do projeto - código se houver]
Data da Reunião: [DD-MMM-AAAA]
Horas: [HH:MM]
Data de criação do Documento: [DD-MMM-AAAA]

Convocados:
- [Nome completo] | [Setor/Departamento] | [Sim ou Não]

Pauta:
- [Item da pauta];

Definições:
- [Definição ou decisão tomada]

Pendências:
- [Descrição da pendência] | [Responsável] | [DD-MMM-AAAA ou A definir]

Plano de Ação:
- [Descrição da ação] | [Responsável] | [DD-MMM-AAAA ou A definir]

IMPORTANTE: Retorne APENAS texto puro no formato acima. Sem negrito, sem Markdown, sem tabelas formatadas, sem numeração automática, sem cabeçalhos decorados. Apenas texto simples com hífens para listas e pipes para separar campos. Não deixe linhas em branco extras dentro de nenhuma seção.
"""


PROMPT_2_TEMPLATE = """Preencha este documento de Ata de Reunião com os dados fornecidos abaixo. O documento contém placeholders em cinza entre colchetes que devem ser substituídos pelos dados reais em texto preto. Siga rigorosamente estas instruções:

INSTRUÇÕES DE PREENCHIMENTO:

1. TABELA "Informações do Projeto":
     - Substitua [Inserir nome do cliente] pelo nome do Cliente.
     - Substitua [Inserir nome e código do projeto] pelo Nome e Código do projeto.
     - Substitua [DD-MMM-AAAA] ao lado de "Data da Reunião" pela data da reunião.
     - Substitua [HH:MM] ao lado de "Horas:" pela hora da reunião. Mantenha o prefixo "Horas: ".
     - Substitua [DD-MMM-AAAA] ao lado de "Data de criação do Documento" pela data do documento.

2. TABELA "Convocados":
     - A tabela tem 1 linha-modelo com placeholders. Para cada pessoa listada nos Convocados, crie uma linha na tabela:
         * Coluna "Convocados": nome completo da pessoa
         * Coluna "Setor e/ou Departamento": setor ou departamento informado
         * Coluna "Ausência": "Sim" ou "Não" conforme informado
     - Se houver mais de 1 pessoa, ADICIONE novas linhas à tabela copiando a formatação da linha-modelo.
     - A linha-modelo original deve ser substituída pelo primeiro participante.

3. SEÇÃO "Pauta":
     - Substitua o placeholder pela lista de itens da pauta.
     - Cada item em uma linha separada.
     - Cada item DEVE terminar com ";" (ponto e vírgula).
     - NÃO deixe linhas em branco entre os itens.

4. SEÇÃO "Definições":
     - Substitua o placeholder pela lista de definições.
     - Cada definição em uma linha separada.
     - As definições NÃO devem ter ";" ao final.
     - NÃO deixe linhas em branco entre os itens nem após o último item.
     - A célula deve se ajustar ao conteúdo, sem espaço vazio sobrando.

5. TABELA "Pendências":
     - Para cada pendência, crie uma linha com:
         * Coluna "Pendências": descrição da pendência
         * Coluna "Responsáveis": nome do responsável
         * Coluna "Previsão": data de previsão (DD-MMM-AAAA) ou "A definir"
     - ADICIONE linhas conforme necessário.

6. TABELA "Plano de Ação":
     - Para cada ação, crie uma linha com:
         * Coluna "Plano de Ação": descrição da ação
         * Coluna "Responsáveis": nome do responsável
         * Coluna "Previsão": data de previsão (DD-MMM-AAAA) ou "A definir"
     - ADICIONE linhas conforme necessário.

REGRAS GERAIS:
- NÃO altere o cabeçalho (header), rodapé (footer), logotipos ou qualquer elemento fixo.
- NÃO altere os títulos das seções ou cabeçalhos das tabelas (linha cinza com texto em negrito).
- MANTENHA toda a formatação original: fontes, tamanhos, cores, bordas, sombreamento.
- O texto preenchido deve estar em preto, não cinza.
- NÃO deixe linhas em branco extras em nenhuma seção.
- Se alguma seção não tiver dados, remova a linha-modelo e deixe apenas o cabeçalho.

DADOS PARA PREENCHIMENTO:

[COLE AQUI O RESULTADO DO PROMPT 1]

Após preencher, salve o arquivo com o nome seguindo este padrão:
[Nome do Cliente com espaços substituídos por _]_[Data da Reunião]_[Hora com h no lugar de :].docx

Exemplo: Alimentos_Naturale_15-Jan-2026_09h00.docx
"""


def build_prompt_2_for_word(preview_text: str) -> str:
        return PROMPT_2_TEMPLATE.replace("[COLE AQUI O RESULTADO DO PROMPT 1]", preview_text.strip())


DEFAULT_MEETING_DATA = {
    "cliente": "Não identificado",
    "nome_codigo_projeto": "Não identificado",
    "data_reuniao": "00-Jan-0000",
    "horas": "00:00",
    "data_documento": "00-Jan-0000",
    "convocados": [],
    "pauta": [],
    "definicoes": [],
    "pendencias": [],
    "plano_acao": [],
}


SECTION_NAMES = {
    "convocados": "Convocados:",
    "pauta": "Pauta:",
    "definicoes": "Definições:",
    "pendencias": "Pendências:",
    "plano_acao": "Plano de ação:",
}


def strip_list_marker(line: str) -> str:
    return re.sub(r"^\s*-\s*", "", line).strip()


def normalize_token(value: str) -> str:
    lowered = unicodedata.normalize("NFKD", value.casefold())
    lowered = "".join(char for char in lowered if not unicodedata.combining(char))
    lowered = re.sub(r"[^a-z0-9]+", " ", lowered)
    return lowered.strip()


def clean_string(value: Any, default: str) -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text or default


def clean_line(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def extract_json_payload(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise ValueError("A resposta do modelo não contém JSON válido.")
        return json.loads(match.group(0))


def iter_block_items(parent: DocumentObject | _Cell):
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def cell_text(cell: _Cell) -> str:
    parts = [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
    return "\n".join(parts).strip()


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.text = text


def apply_data_style(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_text(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        for paragraph in cell.paragraphs:
            apply_data_style(paragraph)
        return
    set_paragraph_text(cell.paragraphs[0], text)
    apply_data_style(cell.paragraphs[0])
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")
        apply_data_style(paragraph)


def set_cell_lines(cell: _Cell, lines: list[str]) -> None:
    lines = [line for line in lines if line.strip()]
    if not lines:
        set_cell_text(cell, "")
        return
    set_paragraph_text(cell.paragraphs[0], lines[0])
    apply_data_style(cell.paragraphs[0])
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")
        apply_data_style(paragraph)
    for line in lines[1:]:
        new_paragraph = cell.add_paragraph(line)
        apply_data_style(new_paragraph)


def remove_row(table: Table, row_index: int) -> None:
    if row_index >= len(table.rows):
        return
    table._tbl.remove(table.rows[row_index]._tr)


def clone_row_from_template(table: Table, template_row_index: int = 1):
    template_row = table.rows[template_row_index]
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def fill_template_by_spec(template_bytes: bytes, preview_text: str) -> tuple[DocumentObject, dict[str, Any]]:
    """Preenche o template seguindo a especificação fixa EAS010-ATA-Template-v3."""
    document = Document(BytesIO(template_bytes))
    payload = preview_to_payload(preview_text)

    if len(document.tables) < 6:
        raise ValueError("Template inválido: esperado pelo menos 6 tabelas no corpo do documento.")

    report = {
        "tabelas_preenchidas": [],
        "linhas_convocados": 0,
        "linhas_pendencias": 0,
        "linhas_plano_acao": 0,
    }

    # Tabela 0 - Informações do Projeto
    t0 = document.tables[0]
    set_cell_text(t0.rows[1].cells[1], payload["cliente"])
    set_cell_text(t0.rows[2].cells[1], payload["nome_codigo_projeto"])
    set_cell_text(t0.rows[3].cells[1], payload["data_reuniao"])
    hora_cell_text = t0.rows[3].cells[2].text
    if "Horas:" in hora_cell_text:
        set_cell_text(t0.rows[3].cells[2], f"Horas: {payload['horas']}")
    else:
        set_cell_text(t0.rows[3].cells[2], payload["horas"])
    set_cell_text(t0.rows[4].cells[1], payload["data_documento"])
    report["tabelas_preenchidas"].append(0)

    # Tabela 1 - Convocados (dinâmica)
    t1 = document.tables[1]
    convocados = payload["convocados"]
    if not convocados:
        remove_row(t1, 1)
    else:
        while len(t1.rows) > 2:
            remove_row(t1, 2)
        set_cell_text(t1.rows[1].cells[0], convocados[0]["nome"])
        set_cell_text(t1.rows[1].cells[1], convocados[0]["setor"])
        set_cell_text(t1.rows[1].cells[2], convocados[0]["ausente"])
        for person in convocados[1:]:
            new_row = clone_row_from_template(t1, 1)
            set_cell_text(new_row.cells[0], person["nome"])
            set_cell_text(new_row.cells[1], person["setor"])
            set_cell_text(new_row.cells[2], person["ausente"])
    report["linhas_convocados"] = len(convocados)
    report["tabelas_preenchidas"].append(1)

    # Tabela 2 - Pauta (multilinha, cada item com ;)
    t2 = document.tables[2]
    pauta_lines = [item if item.endswith(";") else f"{item};" for item in payload["pauta"]]
    set_cell_lines(t2.rows[1].cells[0], pauta_lines)
    report["tabelas_preenchidas"].append(2)

    # Tabela 3 - Definições (multilinha sem ;)
    t3 = document.tables[3]
    definicoes_lines = [item[:-1] if item.endswith(";") else item for item in payload["definicoes"]]
    set_cell_lines(t3.rows[1].cells[0], definicoes_lines)
    report["tabelas_preenchidas"].append(3)

    # Tabela 4 - Pendências (dinâmica)
    t4 = document.tables[4]
    pendencias = payload["pendencias"]
    if not pendencias:
        remove_row(t4, 1)
    else:
        while len(t4.rows) > 2:
            remove_row(t4, 2)
        set_cell_text(t4.rows[1].cells[0], pendencias[0]["descricao"])
        set_cell_text(t4.rows[1].cells[1], pendencias[0]["responsavel"])
        set_cell_text(t4.rows[1].cells[2], pendencias[0]["prazo"])
        for item in pendencias[1:]:
            new_row = clone_row_from_template(t4, 1)
            set_cell_text(new_row.cells[0], item["descricao"])
            set_cell_text(new_row.cells[1], item["responsavel"])
            set_cell_text(new_row.cells[2], item["prazo"])
    report["linhas_pendencias"] = len(pendencias)
    report["tabelas_preenchidas"].append(4)

    # Tabela 5 - Plano de Ação (dinâmica)
    t5 = document.tables[5]
    plano = payload["plano_acao"]
    if not plano:
        remove_row(t5, 1)
    else:
        while len(t5.rows) > 2:
            remove_row(t5, 2)
        set_cell_text(t5.rows[1].cells[0], plano[0]["descricao"])
        set_cell_text(t5.rows[1].cells[1], plano[0]["responsavel"])
        set_cell_text(t5.rows[1].cells[2], plano[0]["prazo"])
        for item in plano[1:]:
            new_row = clone_row_from_template(t5, 1)
            set_cell_text(new_row.cells[0], item["descricao"])
            set_cell_text(new_row.cells[1], item["responsavel"])
            set_cell_text(new_row.cells[2], item["prazo"])
    report["linhas_plano_acao"] = len(plano)
    report["tabelas_preenchidas"].append(5)

    return document, report


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    blocks: list[str] = []
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            text = clean_line(block.text)
            if text:
                blocks.append(text)
            continue
        for row in block.rows:
            row_parts = []
            for cell in row.cells:
                text = clean_line(cell_text(cell))
                if text:
                    row_parts.append(text)
            if row_parts:
                blocks.append(" | ".join(row_parts))
    return "\n".join(blocks)


def extrair_texto_arquivo(uploaded_file: Any) -> str | None:
    if uploaded_file is None:
        return None
    file_name = getattr(uploaded_file, "name", "").lower()
    if file_name.endswith(".docx"):
        return extract_docx_text(uploaded_file.getvalue())
    return uploaded_file.getvalue().decode("utf-8", errors="ignore")


def ensure_people_list(items: list[str]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    for item in items:
        if not item.strip():
            continue
        parts = [part.strip() for part in item.split("|")]
        result.append(
            {
                "nome": parts[0] if len(parts) > 0 and parts[0] else "Não identificado",
                "setor": parts[1] if len(parts) > 1 and parts[1] else "Não identificado",
                "ausente": parts[2] if len(parts) > 2 and parts[2] else "Não",
            }
        )
    return result


def ensure_action_list(items: list[str]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    for item in items:
        if not item.strip():
            continue
        parts = [part.strip() for part in item.split("|")]
        result.append(
            {
                "descricao": parts[0] if len(parts) > 0 and parts[0] else "Não identificado",
                "responsavel": parts[1] if len(parts) > 1 and parts[1] else "A definir",
                "prazo": parts[2] if len(parts) > 2 and parts[2] else "A definir",
            }
        )
    return result


def normalize_meeting_data(data: dict[str, Any] | None) -> dict[str, Any]:
    normalized = deepcopy(DEFAULT_MEETING_DATA)
    if not isinstance(data, dict):
        return normalized

    normalized["cliente"] = clean_string(data.get("cliente"), normalized["cliente"])
    normalized["nome_codigo_projeto"] = clean_string(
        data.get("nome_codigo_projeto", data.get("projeto")),
        normalized["nome_codigo_projeto"],
    )
    normalized["data_reuniao"] = clean_string(data.get("data_reuniao"), normalized["data_reuniao"])
    normalized["horas"] = clean_string(data.get("horas", data.get("hora")), normalized["horas"])
    normalized["data_documento"] = clean_string(data.get("data_documento"), normalized["data_reuniao"])
    normalized["convocados"] = data.get("convocados", []) if isinstance(data.get("convocados"), list) else []
    normalized["pauta"] = data.get("pauta", []) if isinstance(data.get("pauta"), list) else []
    normalized["definicoes"] = data.get("definicoes", []) if isinstance(data.get("definicoes"), list) else []
    normalized["pendencias"] = data.get("pendencias", []) if isinstance(data.get("pendencias"), list) else []
    normalized["plano_acao"] = data.get("plano_acao", []) if isinstance(data.get("plano_acao"), list) else []
    return normalized


def parse_preview_text(preview_text: str) -> dict[str, Any]:
    data = deepcopy(DEFAULT_MEETING_DATA)
    current_section = ""

    for raw_line in preview_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        normalized_line = normalize_token(line)
        if line.startswith("Cliente:"):
            data["cliente"] = clean_string(line.split(":", 1)[1], data["cliente"])
            current_section = ""
            continue
        if line.startswith("Nome e Código do projeto:"):
            data["nome_codigo_projeto"] = clean_string(
                line.split(":", 1)[1],
                data["nome_codigo_projeto"],
            )
            current_section = ""
            continue
        if line.startswith("Data da Reunião:"):
            data["data_reuniao"] = clean_string(line.split(":", 1)[1], data["data_reuniao"])
            current_section = ""
            continue
        if line.startswith("Horas:"):
            data["horas"] = clean_string(line.split(":", 1)[1], data["horas"])
            current_section = ""
            continue
        if line.startswith("Data de criação do Documento:"):
            data["data_documento"] = clean_string(line.split(":", 1)[1], data["data_documento"])
            current_section = ""
            continue

        if normalized_line in {
            "convocados",
            "pauta",
            "definicoes",
            "pendencias",
            "plano de acao",
        }:
            if normalized_line == "plano de acao":
                current_section = "plano_acao"
            else:
                current_section = normalized_line.replace(" ", "_")
            continue

        content = strip_list_marker(line)
        if current_section == "convocados":
            data["convocados"].append(content)
        elif current_section == "pauta":
            data["pauta"].append(content)
        elif current_section == "definicoes":
            data["definicoes"].append(content)
        elif current_section == "pendencias":
            data["pendencias"].append(content)
        elif current_section == "plano_acao":
            data["plano_acao"].append(content)

    return normalize_meeting_data(data)


def compose_preview_text(data: dict[str, Any]) -> str:
    normalized = normalize_meeting_data(data)
    convocados = normalized["convocados"] if normalized["convocados"] else ["Não identificado | Não identificado | Não"]
    pauta = normalized["pauta"] if normalized["pauta"] else ["Não identificado"]
    definicoes = normalized["definicoes"] if normalized["definicoes"] else ["Não identificado"]
    pendencias = normalized["pendencias"] if normalized["pendencias"] else ["Não identificado | A definir | A definir"]
    plano_acao = normalized["plano_acao"] if normalized["plano_acao"] else ["Não identificado | A definir | A definir"]

    sections = [
        f"Cliente: {normalized['cliente']}",
        f"Nome e Código do projeto: {normalized['nome_codigo_projeto']}",
        f"Data da Reunião: {normalized['data_reuniao']}",
        f"Horas: {normalized['horas']}",
        f"Data de criação do Documento: {normalized['data_documento']}",
        "",
        SECTION_NAMES["convocados"],
        *[f"- {item}" for item in convocados],
        "",
        SECTION_NAMES["pauta"],
        *[f"- {item if item.endswith(';') else item + ';'}" for item in pauta],
        "",
        SECTION_NAMES["definicoes"],
        *[f"- {item[:-1] if item.endswith(';') else item}" for item in definicoes],
        "",
        SECTION_NAMES["pendencias"],
        *[f"- {item}" for item in pendencias],
        "",
        SECTION_NAMES["plano_acao"],
        *[f"- {item}" for item in plano_acao],
    ]
    return "\n".join(sections)


def preview_to_payload(preview_text: str) -> dict[str, Any]:
    parsed = parse_preview_text(preview_text)
    return {
        "cliente": parsed["cliente"],
        "nome_codigo_projeto": parsed["nome_codigo_projeto"],
        "data_reuniao": parsed["data_reuniao"],
        "horas": parsed["horas"],
        "data_documento": parsed["data_documento"],
        "convocados": ensure_people_list(parsed["convocados"]),
        "pauta": parsed["pauta"],
        "definicoes": parsed["definicoes"],
        "pendencias": ensure_action_list(parsed["pendencias"]),
        "plano_acao": ensure_action_list(parsed["plano_acao"]),
    }


def find_placeholders(text: str) -> list[str]:
    pattern = re.compile(r"\[([^\]]+)\]|\{\{([^}]+)\}\}")
    found = []
    for match in pattern.finditer(text):
        token = clean_string(match.group(1) or match.group(2), "")
        if token:
            found.append(token)
    return found


def extract_template_blueprint(file_bytes: bytes) -> dict[str, Any]:
    document = Document(BytesIO(file_bytes))
    paragraph_counter = 0
    table_counter = 0
    paragraphs: list[dict[str, str]] = []
    cells: list[dict[str, str]] = []
    blocks: list[dict[str, str]] = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            paragraph_counter += 1
            block_id = f"P{paragraph_counter:03d}"
            text = block.text.strip()
            if text:
                paragraphs.append({"id": block_id, "text": text})
                blocks.append({"id": block_id, "kind": "paragraph", "text": text})
            continue

        table_counter += 1
        for row_index, row in enumerate(block.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = cell_text(cell)
                if not text:
                    continue
                cell_id = f"T{table_counter:03d}R{row_index:03d}C{cell_index:03d}"
                cells.append({"id": cell_id, "text": text})
                blocks.append({"id": cell_id, "kind": "cell", "text": text})

    placeholder_tokens = []
    for block in blocks:
        placeholder_tokens.extend(find_placeholders(block["text"]))

    return {
        "paragraphs": paragraphs,
        "cells": cells,
        "blocks": blocks,
        "placeholders": sorted(set(placeholder_tokens)),
    }


def render_template_overview(blueprint: dict[str, Any]) -> str:
    lines = []
    placeholders = blueprint.get("placeholders", [])
    if placeholders:
        lines.append("Campos detectados no template:")
        lines.extend(placeholders)
    else:
        lines.append("Nenhum placeholder explícito detectado no template.")
    lines.append("")
    lines.append("Estrutura útil do template:")
    for block in blueprint.get("blocks", [])[:30]:
        lines.append(f"{block['id']}: {clean_line(block['text'])}")
    return "\n".join(lines)


def build_prompt_2_input(preview_text: str, blueprint: dict[str, Any]) -> str:
    return json.dumps(
        {
            "pre_ata": preview_text,
            "template_blueprint": blueprint,
        },
        ensure_ascii=False,
    )


def build_document_index(document: DocumentObject) -> dict[str, Paragraph | _Cell]:
    index: dict[str, Paragraph | _Cell] = {}
    paragraph_counter = 0
    table_counter = 0

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            paragraph_counter += 1
            index[f"P{paragraph_counter:03d}"] = block
            continue

        table_counter += 1
        for row_index, row in enumerate(block.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                index[f"T{table_counter:03d}R{row_index:03d}C{cell_index:03d}"] = cell
    return index


def format_people(payload: dict[str, Any]) -> str:
    items = payload.get("convocados", [])
    if not items:
        return "Não identificado | Não identificado | Não"
    return "\n".join(f"{item['nome']} | {item['setor']} | {item['ausente']}" for item in items)


def format_actions(items: list[dict[str, str]]) -> str:
    if not items:
        return "Não identificado | A definir | A definir"
    return "\n".join(f"{item['descricao']} | {item['responsavel']} | {item['prazo']}" for item in items)


def format_lines(items: list[str], default: str = "Não identificado") -> str:
    if not items:
        return default
    return "\n".join(items)


def build_fallback_sections(payload: dict[str, Any]) -> list[dict[str, str]]:
    return [
        {"title": "Cliente", "text": payload.get("cliente", "Não identificado")},
        {"title": "Nome e Código do projeto", "text": payload.get("nome_codigo_projeto", "Não identificado")},
        {"title": "Data da reunião", "text": payload.get("data_reuniao", "00-Jan-0000")},
        {"title": "Horas", "text": payload.get("horas", "00:00")},
        {"title": "Convocados", "text": format_people(payload)},
        {"title": "Pauta", "text": format_lines(payload.get("pauta", []))},
        {"title": "Definições", "text": format_lines(payload.get("definicoes", []))},
        {"title": "Pendências", "text": format_actions(payload.get("pendencias", []))},
        {"title": "Plano de ação", "text": format_actions(payload.get("plano_acao", []))},
    ]


def apply_plan_to_document(template_bytes: bytes, plan: dict[str, Any], preview_text: str) -> tuple[DocumentObject, dict[str, Any]]:
    document = Document(BytesIO(template_bytes))
    index = build_document_index(document)
    payload = preview_to_payload(preview_text)

    applied_updates = 0
    missing_targets: list[str] = []
    for update in plan.get("updates", []):
        if not isinstance(update, dict):
            continue
        target_id = clean_string(update.get("target_id"), "")
        text = clean_string(update.get("text"), "")
        if not target_id or target_id not in index or not text:
            if target_id:
                missing_targets.append(target_id)
            continue
        target = index[target_id]
        if isinstance(target, Paragraph):
            set_paragraph_text(target, text)
        else:
            set_cell_text(target, text)
        applied_updates += 1

    fallback_sections = plan.get("fallback_sections", [])
    if not isinstance(fallback_sections, list):
        fallback_sections = []

    if applied_updates == 0 and not fallback_sections:
        fallback_sections = build_fallback_sections(payload)

    if fallback_sections:
        document.add_page_break()
        document.add_heading("Conteúdo Gerado Automaticamente", level=1)
        for section in fallback_sections:
            title = clean_string(section.get("title"), "Seção")
            text = clean_string(section.get("text"), "Não identificado")
            document.add_heading(title, level=2)
            document.add_paragraph(text)

    report = {
        "updates_aplicados": applied_updates,
        "targets_ausentes": missing_targets,
        "fallback_sections": len(fallback_sections),
        "observacoes_modelo": plan.get("observacoes", []) if isinstance(plan.get("observacoes", []), list) else [],
    }
    return document, report


def extrair_nome_arquivo(cliente: str, data: str, hora: str) -> str:
    safe_cliente = re.sub(r"\s+", "_", clean_string(cliente, "Documento"))
    safe_cliente = re.sub(r"[^\w\-]", "", safe_cliente)
    safe_data = clean_string(data, "00-Jan-0000")
    safe_hora = clean_string(hora, "00:00").replace(":", "h")
    return f"{safe_cliente}_{safe_data}_{safe_hora}.docx"